#!/usr/bin/env python3


# python fill_template.py extraction_structured.json --template template_final_real.docx -o template_gerado.docx

"""Fill a DOCX template using values from extraction_structured.json.

Usage:
  python fill_template.py extraction_structured.json --template template_final_real.docx -o template_gerado.docx

Behavior:
- Reads the structured JSON produced by `analyze_with_openai.py`.
- Replaces placeholders in the DOCX template of the form [Campo] with values found in JSON.
- If `--template` does not exist, a new blank `template_gerado.docx` will be created with placeholders replaced.

Note: this replacement strategy merges runs and may lose complex character-level formatting. It's suitable for templates with plain text placeholders.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from typing import Any, Dict, List, Optional

try:
    from docx import Document
except Exception:
    Document = None


PLACEHOLDER_PATTERN = re.compile(r"\[([^\]]+)\]")


def load_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_nested(data: Dict[str, Any], path: List[str]) -> Optional[Any]:
    cur = data
    for p in path:
        if not isinstance(cur, dict):
            return None
        cur = cur.get(p)
        if cur is None:
            return None
    return cur


def first_non_null(data: Dict[str, Any], paths: List[List[str]]) -> Optional[Any]:
    for p in paths:
        val = get_nested(data, p)
        if val is not None and val != "":
            return val
    return None


def format_paragraph_capitalization(text: str) -> str:
    """Primeira letra minúscula; após ponto final, próxima palavra maiúscula."""
    if not text or not isinstance(text, str):
        return text
    text = text.strip()
    if not text:
        return text
    result = text[0].lower() + text[1:] if len(text) > 1 else text[0].lower()

    def _repl(m) -> str:
        c = m.group(1)
        return ". " + (c.upper() if c.isalpha() else c)

    result = re.sub(r"\.\s+(.)", _repl, result)
    return result


def wrap_long_text(text: str, max_chars: int = 80) -> str:
    """Insere quebras de linha para facilitar a quebra em células/parágrafos do Word."""
    if not text or not isinstance(text, str):
        return text
    text = text.strip()
    if len(text) <= max_chars:
        return text
    lines: List[str] = []
    while text:
        if len(text) <= max_chars:
            lines.append(text)
            break
        break_at = text.rfind(" ", 0, max_chars + 1)
        if break_at <= 0:
            break_at = text.find(" ", max_chars)
        if break_at <= 0:
            lines.append(text)
            break
        lines.append(text[:break_at])
        text = text[break_at + 1 :].lstrip()
    return "\n".join(lines)


def format_date_to_dd_mm_yyyy(val: Optional[Any]) -> str:
    """Converte datas para formato DD/MM/YYYY (ex: 05/12/2024)."""
    if val is None or val == "":
        return ""
    val = str(val).strip()
    if not val:
        return ""
    # ISO YYYY-MM-DD
    m = re.match(r"(\d{4})-(\d{1,2})-(\d{1,2})", val)
    if m:
        y, mo, d = m.group(1), m.group(2).zfill(2), m.group(3).zfill(2)
        return f"{d}/{mo}/{y}"
    # DD/MM/YYYY (já ok)
    if re.match(r"\d{1,2}/\d{1,2}/\d{4}", val):
        parts = val.split("/")
        if len(parts) == 3:
            return f"{parts[0].zfill(2)}/{parts[1].zfill(2)}/{parts[2]}"
    # DD-MM-YYYY
    m = re.match(r"(\d{1,2})-(\d{1,2})-(\d{4})", val)
    if m:
        d, mo, y = m.group(1).zfill(2), m.group(2).zfill(2), m.group(3)
        return f"{d}/{mo}/{y}"
    # "05 de dezembro de 2024"
    _meses = {"janeiro": 1, "fevereiro": 2, "março": 3, "marco": 3, "abril": 4, "maio": 5,
              "junho": 6, "julho": 7, "agosto": 8, "setembro": 9, "outubro": 10,
              "novembro": 11, "dezembro": 12}
    m = re.match(r"(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})", val, re.IGNORECASE)
    if m:
        d, mes_nome, y = m.group(1).zfill(2), m.group(2).lower(), m.group(3)
        mo = _meses.get(mes_nome)
        if mo:
            return f"{d}/{mo:02d}/{y}"
    return val


def prepare_replacements(structured: Dict[str, Any]) -> Dict[str, str]:
    """Map placeholders to values extracted from structured JSON."""
    def s(v: Any) -> str:
        if v is None:
            return ""
        if isinstance(v, (list, dict)):
            return json.dumps(v, ensure_ascii=False)
        return str(v)

    replacements: Dict[str, str] = {}

    # qualification
    replacements["Nome"] = s(first_non_null(structured, [["qualificacao_parte_autora", "nome"]]))
    replacements["nacionalidade"] = s(first_non_null(structured, [["qualificacao_parte_autora", "nacionalidade"]]))
    replacements["estado civil"] = s(first_non_null(structured, [["qualificacao_parte_autora", "estado_civil"]]))
    replacements["profissão"] = s(first_non_null(structured, [["qualificacao_parte_autora", "profissao"]]))
    replacements["CPF"] = s(first_non_null(structured, [["qualificacao_parte_autora", "cpf"]]))
    replacements["Representante legal"] = s(first_non_null(structured, [["qualificacao_parte_autora", "representante_legal_nome"]]))
    replacements["CPF representante"] = s(first_non_null(structured, [["qualificacao_parte_autora", "representante_legal_cpf"]]))
    replacements["RG representante"] = s(first_non_null(structured, [["qualificacao_parte_autora", "representante_legal_rg"]]))
    replacements["endereço completo"] = s(first_non_null(structured, [["qualificacao_parte_autora", "endereco_completo"]]))

    # requerimento INSS
    replacements["Número do benefício"] = s(first_non_null(structured, [["dados_requerimento_inss", "numero_beneficio_NB"]]))
    replacements["data de entrada do requerimento"] = format_date_to_dd_mm_yyyy(first_non_null(structured, [["dados_requerimento_inss", "DER_data_entrada_requerimento"]]))

    # laudos médicos (até 5, ordenados por data)
    laudos = structured.get("dados_medicos", {}).get("laudos", [])
    for i in range(1, 6):  # 1 a 5
        if i <= len(laudos) and laudos[i-1]:
            laudo = laudos[i-1]
            replacements[f"data do laudo médico {i}"] = format_date_to_dd_mm_yyyy(laudo.get("data_laudo"))
            replacements[f"especialidade do médico {i}"] = s(laudo.get("especialidade_medico"))
            replacements[f"nome do médico {i}"] = s(laudo.get("nome_medico"))
            
            # Descrição com verbo no início
            desc = s(laudo.get("descricao", ""))
            if desc:
                desc = desc.strip()
                # Garante que começa com verbo em minúscula
                verbos_comuns = ["recomenda", "atesta", "conclui", "indica", "sugere", "prescreve"]
                starts_with_verb = any(desc.lower().startswith(v) for v in verbos_comuns)
                if not starts_with_verb:
                    desc = "recomenda " + desc
                replacements[f"descrição do laudo {i}"] = format_paragraph_capitalization(desc)
            else:
                replacements[f"descrição do laudo {i}"] = ""
        else:
            # Laudo não existe, deixa em branco
            replacements[f"data do laudo médico {i}"] = ""
            replacements[f"especialidade do médico {i}"] = ""
            replacements[f"nome do médico {i}"] = ""
            replacements[f"descrição do laudo {i}"] = ""

    # relatorio escolar
    replacements["data de emissão do relatório escolar"] = format_date_to_dd_mm_yyyy(first_non_null(structured, [["dados_medicos", "relatorio_escolar", "data_emissao"]]))
    replacements["primeiro nome do autor"] = s(first_non_null(structured, [["dados_medicos", "relatorio_escolar", "primeiro_nome_do_autor"]]))
    resumo_esc = s(first_non_null(structured, [["dados_medicos", "relatorio_escolar", "resumo"]]))
    replacements["resumo do relatório escolar"] = format_paragraph_capitalization(resumo_esc) if resumo_esc else ""
    resumo_cont = s(first_non_null(structured, [["dados_medicos", "relatorio_escolar", "resumo_continuacao"]]))
    replacements["continuação do resumo do relatório escolar"] = format_paragraph_capitalization(resumo_cont) if resumo_cont else ""
    replacements["continua o resumo do relatório escolar"] = replacements["continuação do resumo do relatório escolar"]

    # segundo laudo
    replacements["data do “segundo laudo”"] = format_date_to_dd_mm_yyyy(first_non_null(structured, [["dados_medicos", "laudo_psiquiatrico_segundo_laudo", "data_segundo_laudo"]]))
    replacements["Nome do médico segundo"] = s(first_non_null(structured, [["dados_medicos", "laudo_psiquiatrico_segundo_laudo", "nome_medico"]]))
    replacements["Nome do médico"] = replacements["Nome do médico segundo"]  # alias (laudo psiquiátrico)
    resumo_laudo = s(first_non_null(structured, [["dados_medicos", "laudo_psiquiatrico_segundo_laudo", "resumo"]]))
    replacements["resumo do laudo médico"] = format_paragraph_capitalization(resumo_laudo) if resumo_laudo else ""


    # diagnostico e tratamento
    conclusao = s(first_non_null(structured, [["dados_medicos", "diagnostico_final_tratamento", "conclusao_medica"]]))
    replacements["conclusão médica"] = format_paragraph_capitalization(conclusao) if conclusao else ""
    def_cid = s(first_non_null(structured, [["dados_medicos", "diagnostico_final_tratamento", "deficiencia_e_CID"]]))
    replacements["deficiência e a CID correspondente"] = format_paragraph_capitalization(wrap_long_text(def_cid)) if def_cid else ""
    def_assoc = s(first_non_null(structured, [["dados_medicos", "diagnostico_final_tratamento", "deficiencia_associada_e_CID"]]))
    replacements["deficiência associada e CID correspondente"] = format_paragraph_capitalization(wrap_long_text(def_assoc)) if def_assoc else ""
    replacements["medicamento prescrito"] = s(first_non_null(structured, [["dados_medicos", "diagnostico_final_tratamento", "medicamento_prescrito"]]))
    finalidade = s(first_non_null(structured, [["dados_medicos", "diagnostico_final_tratamento", "finalidade_medicamento"]]))
    replacements["descrição da finalidade do medicamento"] = format_paragraph_capitalization(finalidade) if finalidade else ""
    replacements["descrição para que serve o medicamento"] = replacements["descrição da finalidade do medicamento"]  # alias

    # dados socioeconomicos
    replacements["detalhar o grau de parentesco das pessoas listadas no cadastro único"] = s(first_non_null(structured, [["dados_socioeconomicos", "grau_parentesco_CadUnico"]]))
    replacements["nome da avó"] = s(first_non_null(structured, [["dados_socioeconomicos", "nome_avo"]]))
    replacements["valor exato da aposentadoria – R$ ___ "] = s(first_non_null(structured, [["dados_socioeconomicos", "valor_exato_aposentadoria"]]))
    replacements["páginas do laudo social – “anexo 05, pgs. XX”"] = s(first_non_null(structured, [["dados_socioeconomicos", "paginas_laudo_social"]]))

    # dados processuais
    replacements["Número do benefício / NB"] = replacements.get("Número do benefício", "")

    return replacements


def _replace_across_runs(paragraph, placeholder: str, value: str) -> bool:
    """Trata casos onde o placeholder está dividido entre múltiplos runs.
    
    Retorna True se encontrou e substituiu o placeholder.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    search_pattern = f"[{placeholder}]"
    
    if search_pattern not in full_text:
        return False
    
    # Encontra a posição do placeholder no texto completo
    start_pos = full_text.index(search_pattern)
    end_pos = start_pos + len(search_pattern)
    
    # Mapeia posições para runs
    current_pos = 0
    runs_to_modify = []
    
    for i, run in enumerate(paragraph.runs):
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        # Este run contém parte do placeholder?
        if run_start < end_pos and run_end > start_pos:
            overlap_start = max(0, start_pos - run_start)
            overlap_end = min(len(run.text), end_pos - run_start)
            runs_to_modify.append((i, run, overlap_start, overlap_end, run_start))
        
        current_pos = run_end
    
    if not runs_to_modify:
        return False
    
    # Substitui o placeholder mantendo formatação do primeiro run afetado
    first_idx, first_run, first_start, first_end, first_abs_start = runs_to_modify[0]
    
    # Reconstrói o texto do primeiro run
    before = first_run.text[:first_start]
    after = first_run.text[first_end:]
    first_run.text = before + value + after
    
    # Remove o texto dos runs subsequentes que faziam parte do placeholder
    for idx, run, overlap_start, overlap_end, abs_start in runs_to_modify[1:]:
        before = run.text[:overlap_start]
        after = run.text[overlap_end:]
        run.text = before + after
    
    return True


def replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """Substitui placeholders preservando a formatação (negrito, itálico, etc.) dos runs.
    
    Esta função trata dois casos:
    1. Placeholder completo dentro de um único run (caso simples)
    2. Placeholder dividido entre múltiplos runs (caso complexo)
    """
    if not paragraph.runs:
        return
    
    # Primeiro passo: tenta substituir em cada run individualmente (caso simples)
    for run in paragraph.runs:
        if not run.text:
            continue
        original_text = run.text
        for placeholder, value in replacements.items():
            if f"[{placeholder}]" in run.text:
                run.text = run.text.replace(f"[{placeholder}]", value)
    
    # Segundo passo: verifica se ainda há placeholders não substituídos
    # (isso indica que estão divididos entre runs)
    full_text = "".join(run.text for run in paragraph.runs)
    
    for placeholder, value in replacements.items():
        search_pattern = f"[{placeholder}]"
        if search_pattern in full_text:
            # Placeholder está dividido entre runs - usa função auxiliar
            _replace_across_runs(paragraph, placeholder, value)
            # Atualiza o texto completo após a substituição
            full_text = "".join(run.text for run in paragraph.runs)


def replace_in_table(table: Any, replacements: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, replacements)


def process_document(doc: Any, replacements: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
    for tbl in doc.tables:
        replace_in_table(tbl, replacements)


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Fill DOCX template with fields from structured JSON")
    parser.add_argument("input_json", help="Structured JSON file (from analyze_with_openai.py)")
    parser.add_argument("--template", default="template_final_real.docx", help="Path to DOCX template")
    parser.add_argument("-o", "--output", default="template_gerado.docx", help="Output filled DOCX")
    args = parser.parse_args(argv)

    if Document is None:
        print("Missing dependency 'python-docx'. Install with: pip install python-docx", file=sys.stderr)
        return 2

    if not os.path.exists(args.input_json):
        print(f"Input JSON not found: {args.input_json}", file=sys.stderr)
        return 2

    structured = load_json(args.input_json)
    replacements = prepare_replacements(structured)

    # If template not exists, create a blank doc with placeholders inserted
    if not os.path.exists(args.template):
        doc = Document()
        # create a minimal template based on the user's sample header
        doc.add_paragraph("Ao Juízo da 27° Vara do Juizado Especial Federal da Comarca de Itapipoca/CE.")
        doc.add_paragraph("[Nome], [nacionalidade], [estado civil], [profissão], [CPF:], representado (a) por [Representante legal], [CPF representante] e [RG representante], residentes e domiciliados (as) na [endereço completo] (anexo 01), ...")
        doc.save(args.output)
        print(f"Template not found. Created minimal template and saved to {args.output}")
        return 0

    doc = Document(args.template)
    process_document(doc, replacements)
    doc.save(args.output)
    print(f"Generated filled document: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
