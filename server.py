#!/usr/bin/env python3
from __future__ import annotations

import io
import os
import json
import re
import time
import uuid
import tempfile
import threading
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from enum import Enum

import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from fastapi import Body, Depends, FastAPI, File, Header, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

from docx import Document
import firebase_admin
from firebase_admin import auth as firebase_auth
from firebase_admin import credentials as firebase_credentials

from analyze_with_openai import (
    DEFAULT_MODEL,
    FUNCTION_SCHEMA,
    build_text_from_ocr,
    call_openai_chat,
    extract_from_response,
    load_env,
    merge_structured_field,
    select_page_numbers_for_extraction,
)
from fill_template import prepare_replacements, process_document


APP_ROOT = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(APP_ROOT, "static")
TEMPLATE_PATH = os.path.join(APP_ROOT, "template_final_real.docx")
LOGS_DIR = os.path.join(APP_ROOT, "logs")
INSS_LOG_FILE = os.path.join(LOGS_DIR, "inss_extraction.log")

# armazenamento temporário em memória (token -> payload)
_STORE: Dict[str, Dict[str, Any]] = {}
_STORE_TTL_SECONDS = 60 * 30  # 30 minutos
_ALWAYS_NACIONALIDADE = "brasileiro(a)"

_FIREBASE_APP: Optional[firebase_admin.App] = None

# Sistema de tarefas em background
class TaskStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"

class Task:
    def __init__(self, task_id: str):
        self.task_id = task_id
        self.status = TaskStatus.PENDING
        self.progress = 0
        self.message = "Tarefa criada"
        self.result = None
        self.error = None
        self.created_at = time.time()
        self.updated_at = time.time()
    
    def update(self, status: Optional[TaskStatus] = None, progress: Optional[int] = None, 
               message: Optional[str] = None, result: Optional[Any] = None, error: Optional[str] = None):
        if status is not None:
            self.status = status
        if progress is not None:
            self.progress = progress
        if message is not None:
            self.message = message
        if result is not None:
            self.result = result
        if error is not None:
            self.error = error
        self.updated_at = time.time()
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "task_id": self.task_id,
            "status": self.status,
            "progress": self.progress,
            "message": self.message,
            "result": self.result,
            "error": self.error,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
        }

_TASKS: Dict[str, Task] = {}
_TASKS_LOCK = threading.Lock()


def _init_firebase_admin() -> firebase_admin.App:
    """Inicializa firebase-admin (lazy) para validação de idToken."""
    global _FIREBASE_APP
    if _FIREBASE_APP is not None:
        return _FIREBASE_APP

    # 1) Path para serviceAccountKey.json
    sa_path = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON")
    if sa_path:
        cred = firebase_credentials.Certificate(sa_path)
        _FIREBASE_APP = firebase_admin.initialize_app(cred, {"projectId": os.environ.get("FIREBASE_PROJECT_ID")})
        return _FIREBASE_APP

    # 2) JSON string do service account
    sa_json = os.environ.get("FIREBASE_SERVICE_ACCOUNT")
    if sa_json:
        info = json.loads(sa_json)
        cred = firebase_credentials.Certificate(info)
        _FIREBASE_APP = firebase_admin.initialize_app(cred, {"projectId": os.environ.get("FIREBASE_PROJECT_ID")})
        return _FIREBASE_APP

    # 3) GOOGLE_APPLICATION_CREDENTIALS (ADC) ou credenciais padrão
    options: Dict[str, Any] = {}
    project_id = os.environ.get("FIREBASE_PROJECT_ID")
    if project_id:
        options["projectId"] = project_id
    _FIREBASE_APP = firebase_admin.initialize_app(options=options or None)
    return _FIREBASE_APP


def require_firebase_user(authorization: Optional[str] = Header(default=None)) -> Dict[str, Any]:
    """Dependency do FastAPI: exige Authorization: Bearer <idToken> e valida com firebase-admin."""
    if not authorization or not isinstance(authorization, str):
        print("[AUTH] Cabeçalho Authorization ausente ou inválido")
        raise HTTPException(status_code=401, detail="Não autenticado. Envie Authorization: Bearer <idToken>.")
    if not authorization.lower().startswith("bearer "):
        print(f"[AUTH] Cabeçalho Authorization não começa com 'Bearer ': {authorization[:20]}...")
        raise HTTPException(status_code=401, detail="Cabeçalho Authorization inválido. Use Bearer <idToken>.")

    token = authorization.split(" ", 1)[1].strip()
    if not token:
        print("[AUTH] Token vazio após split")
        raise HTTPException(status_code=401, detail="Token ausente. Use Authorization: Bearer <idToken>.")

    try:
        _init_firebase_admin()
        check_revoked = os.environ.get("FIREBASE_CHECK_REVOKED", "").strip() in ("1", "true", "yes", "on")
        decoded = firebase_auth.verify_id_token(token, check_revoked=check_revoked)
        if not isinstance(decoded, dict) or not decoded.get("uid"):
            print("[AUTH] Token decodificado mas sem uid válido")
            raise ValueError("Decoded token inválido.")
        print(f"[AUTH] Token válido para usuário: {decoded.get('email', decoded.get('uid'))}")
        return decoded
    except Exception as e:
        # Log do erro real (para debug do servidor, não para o cliente)
        print(f"[AUTH] Erro ao validar token: {type(e).__name__}: {str(e)}")
        # não vaza detalhes de validação/credencial para o cliente
        raise HTTPException(status_code=401, detail="Sessão inválida ou expirada. Faça login novamente.")


def _cleanup_store() -> None:
    now = time.time()
    expired = [k for k, v in _STORE.items() if (now - float(v.get("created_at", now))) > _STORE_TTL_SECONDS]
    for k in expired:
        _STORE.pop(k, None)

def _cleanup_tasks() -> None:
    """Remove tarefas antigas (mais de 1 hora)."""
    now = time.time()
    with _TASKS_LOCK:
        expired = [k for k, task in _TASKS.items() if (now - task.created_at) > 3600]
        for k in expired:
            _TASKS.pop(k, None)


def _force_brazilian_nationality(structured: Dict[str, Any]) -> Dict[str, Any]:
    """Força a nacionalidade como 'brasileiro(a)' no JSON estruturado."""
    try:
        qa = structured.get("qualificacao_parte_autora")
        if not isinstance(qa, dict):
            qa = {}
            structured["qualificacao_parte_autora"] = qa
        qa["nacionalidade"] = _ALWAYS_NACIONALIDADE
    except Exception:
        # best-effort
        return structured
    return structured


def _capitalize_name(name: str) -> str:
    """Capitaliza nomes próprios corretamente (primeira letra de cada palavra maiúscula)."""
    if not name or not isinstance(name, str):
        return name
    
    # Lista de preposições e artigos que devem ficar em minúscula (exceto se forem a primeira palavra)
    lowercase_words = {'da', 'de', 'do', 'das', 'dos', 'e', 'em', 'na', 'no', 'a', 'o'}
    
    words = name.strip().split()
    result = []
    
    for i, word in enumerate(words):
        # Primeira palavra sempre capitalizada
        if i == 0:
            result.append(word.capitalize())
        # Preposições e artigos em minúscula
        elif word.lower() in lowercase_words:
            result.append(word.lower())
        # Outras palavras capitalizadas
        else:
            result.append(word.capitalize())
    
    return ' '.join(result)


def _format_cpf(cpf: str) -> str:
    """Formata CPF para XXX.XXX.XXX-XX."""
    if not cpf or not isinstance(cpf, str):
        return cpf
    
    # Remove tudo que não é dígito
    digits = ''.join(c for c in cpf if c.isdigit())
    
    # Se não tem 11 dígitos, retorna original
    if len(digits) != 11:
        return cpf
    
    # Formata XXX.XXX.XXX-XX
    return f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"


_MESES_PT = {
    "janeiro": 1, "fevereiro": 2, "março": 3, "marco": 3, "abril": 4, "maio": 5,
    "junho": 6, "julho": 7, "agosto": 8, "setembro": 9, "outubro": 10,
    "novembro": 11, "dezembro": 12,
}


def _format_date(val: str) -> str:
    """Converte datas para formato DD/MM/YYYY (ex: 05/12/2024)."""
    if not val or not isinstance(val, str):
        return val
    val = val.strip()
    if not val:
        return val
    # ISO YYYY-MM-DD
    m = re.match(r"(\d{4})-(\d{1,2})-(\d{1,2})", val)
    if m:
        y, mo, d = m.group(1), m.group(2).zfill(2), m.group(3).zfill(2)
        return f"{d}/{mo}/{y}"
    # DD/MM/YYYY (já está ok)
    if re.match(r"\d{1,2}/\d{1,2}/\d{4}", val):
        parts = val.split("/")
        if len(parts) == 3:
            return f"{parts[0].zfill(2)}/{parts[1].zfill(2)}/{parts[2]}"
    # DD-MM-YYYY
    m = re.match(r"(\d{1,2})-(\d{1,2})-(\d{4})", val)
    if m:
        d, mo, y = m.group(1).zfill(2), m.group(2).zfill(2), m.group(3)
        return f"{d}/{mo}/{y}"
    # "05 de dezembro de 2024" ou "5 de dezembro de 2024"
    m = re.match(r"(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})", val, re.IGNORECASE)
    if m:
        d, mes_nome, y = m.group(1).zfill(2), m.group(2).lower(), m.group(3)
        mo = _MESES_PT.get(mes_nome)
        if mo:
            return f"{d}/{mo:02d}/{y}"
    return val


def _format_cid(cid: str) -> str:
    """Formata CID para garantir formato CID-XX ou CID-XX.X."""
    if not cid or not isinstance(cid, str):
        return cid
    
    cid = cid.strip().upper()
    
    # Se já começa com CID, retorna normalizado
    if cid.startswith('CID'):
        # Remove espaços extras e normaliza
        cid = cid.replace('CID', 'CID-').replace('--', '-').replace(' ', '')
        return cid
    
    # Se é só o código (ex: F84.0), adiciona CID-
    if len(cid) >= 3 and cid[0].isalpha():
        return f"CID-{cid}"
    
    return cid


def _capitalize_medical_term(term: str) -> str:
    """Capitaliza termos médicos/especialidades (primeira letra de cada palavra importante)."""
    if not term or not isinstance(term, str):
        return term
    
    # Lista de palavras que devem ficar em minúscula
    lowercase_words = {'de', 'da', 'do', 'das', 'dos', 'e', 'em', 'na', 'no', 'a', 'o', 'com', 'por'}
    
    # Exceções: siglas médicas que devem ficar maiúsculas
    uppercase_terms = {'tea', 'tdah', 'toc', 'tpt', 'tag'}
    
    words = term.strip().split()
    result = []
    
    for i, word in enumerate(words):
        word_lower = word.lower()
        
        # Primeira palavra sempre capitalizada
        if i == 0:
            if word_lower in uppercase_terms:
                result.append(word.upper())
            else:
                result.append(word.capitalize())
        # Siglas em maiúscula
        elif word_lower in uppercase_terms:
            result.append(word.upper())
        # Preposições em minúscula
        elif word_lower in lowercase_words:
            result.append(word_lower)
        # Outras palavras capitalizadas
        else:
            result.append(word.capitalize())
    
    return ' '.join(result)


def _get_field_type(path: List[str]) -> str:
    """Determina o tipo de formatação necessária para um campo."""
    if not path:
        return 'lowercase'
    
    last = path[-1]
    
    # Nomes próprios (pessoas)
    name_fields = {
        'nome',
        'representante_legal_nome',
        'nome_do_medico',
        'nome_medico',
        'primeiro_nome_do_autor',
        'nome_avo',
    }
    
    # Endereços
    address_fields = {
        'endereco_completo',
    }
    
    # CPF/RG
    cpf_fields = {
        'cpf',
        'representante_legal_cpf',
    }
    
    # CID
    cid_fields = {
        'CID_da_doenca',
        'cid',
    }
    
    # Datas (formato DD/MM/YYYY)
    date_fields = {
        'data_emissao',
        'data_do_laudo',
        'data_segundo_laudo',
        'DER_data_entrada_requerimento',
    }
    
    # Especialidades médicas, deficiências, medicamentos
    medical_fields = {
        'especialidade_do_medico',
        'deficiencia_constatada',
        'deficiencia_e_CID',
        'deficiencia_associada_e_CID',
        'medicamento_prescrito',
    }
    
    if last in name_fields or last in address_fields:
        return 'name'
    elif last in cpf_fields:
        return 'cpf'
    elif last in cid_fields:
        return 'cid'
    elif last in date_fields:
        return 'date'
    elif last in medical_fields:
        return 'medical'
    else:
        return 'lowercase'


def _process_strings(obj: Any, path: List[str] = None) -> Any:
    """Processa strings recursivamente: formata nomes, CPF, CID, termos médicos, etc."""
    if path is None:
        path = []
    
    if isinstance(obj, str):
        field_type = _get_field_type(path)
        
        if field_type == 'name':
            return _capitalize_name(obj)
        elif field_type == 'cpf':
            return _format_cpf(obj)
        elif field_type == 'cid':
            return _format_cid(obj)
        elif field_type == 'date':
            return _format_date(obj)
        elif field_type == 'medical':
            return _capitalize_medical_term(obj)
        else:
            return obj.lower()
    
    if isinstance(obj, list):
        return [_process_strings(v, path) for v in obj]
    
    if isinstance(obj, dict):
        return {k: _process_strings(v, path + [k]) for k, v in obj.items()}
    
    return obj

def _log_inss_extraction(source: str, chunk_idx: int, total_chunks: int, text_sample: str, structured: Dict[str, Any]) -> None:
    """Registra em log a extração do documento INSS."""
    try:
        os.makedirs(LOGS_DIR, exist_ok=True)
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        text_preview = (text_sample or "")[:500].replace("\n", " ")
        entry = {
            "timestamp": ts,
            "source": source,
            "chunk": f"{chunk_idx + 1}/{total_chunks}",
            "text_preview": text_preview,
            "structured": structured,
        }
        with open(INSS_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"[LOG INSS] Erro ao gravar log: {e}")


def _configure_ocr() -> Dict[str, Optional[str]]:
    """Configura Tesseract/Poppler via variáveis de ambiente (Windows-friendly).

    - TESSERACT_CMD: caminho do tesseract.exe
    - POPPLER_PATH: pasta do poppler/bin (para pdf2image)
    """
    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if tesseract_cmd:
        # pytesseract usa esse atributo para apontar para o executável
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    poppler_path = os.environ.get("POPPLER_PATH")
    return {"tesseract_cmd": tesseract_cmd, "poppler_path": poppler_path}


def _should_fallback_to_ocr(text: str, min_chars: int) -> bool:
    # heurística simples: se o texto for muito curto, é provável que a página seja imagem/scan
    return len((text or "").strip()) < min_chars


def _ocr_page_from_pdf_bytes(
    file_bytes: bytes,
    page_number_1based: int,
    lang: str,
    dpi: int,
    poppler_path: Optional[str],
) -> str:
    images = convert_from_bytes(
        file_bytes,
        dpi=dpi,
        first_page=page_number_1based,
        last_page=page_number_1based,
        poppler_path=poppler_path,
    )
    if not images:
        return ""
    # config: PSM 6 costuma funcionar bem para páginas com blocos de texto
    return pytesseract.image_to_string(images[0], lang=lang, config="--psm 6")


def extract_pdf_hybrid_text(
    file_bytes: bytes,
    source_name: str,
    *,
    ocr_lang: str = "por",
    ocr_min_chars: int = 30,
    ocr_dpi: int = 300,
) -> Dict[str, Any]:
    """Extrai texto do PDF (pdfplumber) com fallback para OCR (Tesseract) por página."""
    cfg = _configure_ocr()
    pages: List[Dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            method = "pdfplumber"
            if _should_fallback_to_ocr(text, ocr_min_chars):
                try:
                    ocr_text = _ocr_page_from_pdf_bytes(
                        file_bytes,
                        page_number_1based=i,
                        lang=ocr_lang,
                        dpi=ocr_dpi,
                        poppler_path=cfg.get("poppler_path"),
                    )
                    if ocr_text and len(ocr_text.strip()) >= len(text.strip()):
                        text = ocr_text
                        method = "ocr"
                except Exception:
                    # se OCR falhar, mantemos o texto original (mesmo que vazio)
                    pass

            # words não são usados no pipeline atual; manter vazio reduz payload/tempo
            pages.append({"page_number": i, "text": text, "words": [], "method": method})
    return {"source": source_name, "page_count": len(pages), "pages": pages}


def extract_image_text(
    image_bytes: bytes,
    source_name: str,
    *,
    ocr_lang: str = "por",
    ocr_dpi: int = 300,
) -> Dict[str, Any]:
    """Extrai texto de imagem PNG/JPG usando Tesseract OCR."""
    from PIL import Image
    
    _configure_ocr()
    
    # Carrega a imagem
    img = Image.open(io.BytesIO(image_bytes))
    
    # Aplica OCR
    text = pytesseract.image_to_string(img, lang=ocr_lang, config="--psm 6")
    
    # Retorna no mesmo formato que PDF (1 página)
    pages = [{"page_number": 1, "text": text, "words": [], "method": "ocr"}]
    return {"source": source_name, "page_count": 1, "pages": pages}


def _summarize_structured(structured: Dict[str, Any]) -> Dict[str, Any]:
    """Gera um resumo “human-friendly” do JSON estruturado para o modal."""
    def g(path: List[str]) -> Optional[Any]:
        cur: Any = structured
        for p in path:
            if not isinstance(cur, dict):
                return None
            cur = cur.get(p)
            if cur is None:
                return None
        return cur

    return {
        "qualificacao": {
            "nome": g(["qualificacao_parte_autora", "nome"]),
            "cpf": g(["qualificacao_parte_autora", "cpf"]),
            "endereco": g(["qualificacao_parte_autora", "endereco_completo"]),
            "representante_legal": g(["qualificacao_parte_autora", "representante_legal_nome"]),
        },
        "inss": {
            "nb": g(["dados_requerimento_inss", "numero_beneficio_NB"]) or g(["dados_processuais", "numero_beneficio_NB_repetido"]),
            "der": g(["dados_requerimento_inss", "DER_data_entrada_requerimento"]),
        },
        "laudo_principal": {
            "deficiencia": g(["dados_medicos", "laudo_principal", "deficiencia_constatada"]),
            "cid": g(["dados_medicos", "laudo_principal", "CID_da_doenca"]),
            "data": g(["dados_medicos", "laudo_principal", "data_do_laudo"]),
            "medico": g(["dados_medicos", "laudo_principal", "nome_do_medico"]),
        },
        "relatorio_escolar": {
            "data_emissao": g(["dados_medicos", "relatorio_escolar", "data_emissao"]),
            "primeiro_nome_autor": g(["dados_medicos", "relatorio_escolar", "primeiro_nome_do_autor"]),
        },
    }


def _process_extraction_background(
    task_id: str,
    file_contents: List[Tuple[bytes, str]],
    model: str,
    pages_per_call: int,
    delay_between_calls: float,
    inss_max_pages: int,
    inss_min_total_pages: int,
    ocr_lang: str,
    ocr_min_chars: int,
    ocr_dpi: int,
) -> None:
    """Processa a extração em background e atualiza o status da tarefa."""
    task = _TASKS.get(task_id)
    if not task:
        return
    
    try:
        task.update(status=TaskStatus.PROCESSING, progress=10, message="Extraindo texto dos arquivos...")
        
        # Extrai texto dos arquivos (PDF, PNG, JPG)
        extractions: List[Dict[str, Any]] = []
        for i, (content, filename) in enumerate(file_contents):
            task.update(progress=10 + (i * 20 // len(file_contents)), 
                       message=f"Extraindo texto de {filename}...")
            
            # Detecta o tipo de arquivo pela extensão
            ext = filename.lower()
            
            if ext.endswith('.pdf'):
                extraction = extract_pdf_hybrid_text(
                    content,
                    filename,
                    ocr_lang=ocr_lang,
                    ocr_min_chars=ocr_min_chars,
                    ocr_dpi=ocr_dpi,
                )
            elif ext.endswith(('.png', '.jpg', '.jpeg')):
                extraction = extract_image_text(
                    content,
                    filename,
                    ocr_lang=ocr_lang,
                    ocr_dpi=ocr_dpi,
                )
            else:
                raise ValueError(f"Tipo de arquivo não suportado: {filename}")
            
            extractions.append(extraction)
        
        task.update(progress=30, message="Analisando documentos com IA...")
        
        # Analisa com OpenAI
        structured, per_doc = analyze_extractions_with_openai(
            extractions=extractions,
            model=model,
            pages_per_call=pages_per_call,
            delay_between_calls=delay_between_calls,
            inss_max_pages=inss_max_pages,
            inss_min_total_pages=inss_min_total_pages,
            task=task,  # passa a task para atualizar progresso
        )
        
        if not structured:
            task.update(status=TaskStatus.FAILED, progress=100, 
                       error="A API não retornou nenhum conteúdo estruturado.")
            return
        
        task.update(progress=95, message="Finalizando...")
        
        structured = _force_brazilian_nationality(structured)
        structured = _process_strings(structured)
        
        token = str(uuid.uuid4())
        _STORE[token] = {"created_at": time.time(), "structured": structured}
        
        result = {
            "token": token,
            "per_document": per_doc,
            "summary": _summarize_structured(structured),
            "structured": structured,
        }
        
        task.update(status=TaskStatus.COMPLETED, progress=100, 
                   message="Extração concluída com sucesso!", result=result)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[TASK {task_id}] Erro ao processar: {type(e).__name__}: {str(e)}")
        print(f"[TASK {task_id}] Traceback completo:\n{error_details}")
        
        # Tenta extrair mais detalhes do erro OpenAI
        error_message = str(e)
        if hasattr(e, '__cause__') and e.__cause__:
            error_message += f" | Causa: {str(e.__cause__)}"
        if hasattr(e, 'response') and e.response:
            error_message += f" | Response: {e.response}"
        
        task.update(status=TaskStatus.FAILED, progress=100, 
                   error=error_message, message="Erro ao processar documentos")

def analyze_extractions_with_openai(
    extractions: List[Dict[str, Any]],
    model: str,
    pages_per_call: int,
    delay_between_calls: float,
    inss_max_pages: int,
    inss_min_total_pages: int,
    task: Optional[Task] = None,
) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """Roda a etapa de estruturação (OpenAI) por documento, retornando resultado + metadados por doc."""
    system_msg = (
        "Você é um assistente que analisa laudos e extrai campos jurídicos e médicos. "
        "Retorne apenas JSON seguindo o schema de função. Se o dado não existir, retorne null. "
        "Datas: use formato DD/MM/YYYY (ex: 05/12/2024). "
        "Se houver baixa confiança, inclua o valor e marque '[confiança baixa]'."
    )

    structured_result: Dict[str, Any] = {}
    per_doc: List[Dict[str, Any]] = []

    INSS_KEYWORDS = ["inss", "instituto nacional", "do seguro social"]

    def _is_inss_doc(ex: Dict[str, Any]) -> bool:
        # Documento do INSS: detecta por palavras-chave na página 1.
        # Não depende do total de páginas, para garantir que analisaremos as 6 primeiras páginas
        # mesmo em PDFs menores/recortados.
        t1 = ""
        if isinstance(ex.get("pages"), list) and ex["pages"]:
            t1 = (ex["pages"][0].get("text") or "")
        t1 = t1.lower()
        return any(k in t1 for k in INSS_KEYWORDS)

    def _is_laudo_doc(ex: Dict[str, Any]) -> bool:
        t1 = ""
        if isinstance(ex.get("pages"), list) and ex["pages"]:
            t1 = (ex["pages"][0].get("text") or "").lower()
        return ("laudo" in t1) and ("médic" in t1 or "medic" in t1 or "crm" in t1 or "rqe" in t1)

    def _score_psiquiatria(ex: Dict[str, Any], pages: List[int]) -> int:
        score = 0
        needles = ["psiquiatr", "psiquiátr", "psiquiatria", "psiquiatra"]
        # olha só nas duas primeiras páginas analisadas
        for pn in pages[: min(len(pages), 2)]:
            txt = ""
            for p in ex.get("pages", []) or []:
                if p.get("page_number") == pn:
                    txt = (p.get("text") or "").lower()
                    break
            for n in needles:
                score += txt.count(n)
        return score

    # pré-calcula páginas analisadas e papéis dos laudos (principal x segundo)
    pages_by_source: Dict[str, List[int]] = {}
    laudo_candidates: List[Tuple[str, int]] = []
    for ex_idx, ex in enumerate(extractions, start=1):
        source = ex.get("source") or f"documento_{ex_idx}"
        pages = select_page_numbers_for_extraction(
            ex,
            selected_pages=None,
            inss_max_pages=inss_max_pages,
            inss_min_total_pages=inss_min_total_pages,
        )
        pages_by_source[source] = pages
        if _is_laudo_doc(ex):
            laudo_candidates.append((source, _score_psiquiatria(ex, pages)))

    laudo_roles: Dict[str, str] = { (ex.get("source") or f"documento_{i+1}") : "none" for i, ex in enumerate(extractions) }
    if len(laudo_candidates) >= 2:
        segundo = sorted(enumerate(laudo_candidates), key=lambda it: (-it[1][1], it[0]))[0][1][0]
        principal = next((s for (s, _) in laudo_candidates if s != segundo), laudo_candidates[0][0])
        laudo_roles[principal] = "principal"
        laudo_roles[segundo] = "segundo"
    elif len(laudo_candidates) == 1:
        laudo_roles[laudo_candidates[0][0]] = "principal"

    for ex_idx, ex in enumerate(extractions, start=1):
        source = ex.get("source") or f"documento_{ex_idx}"
        page_numbers = pages_by_source.get(source) or []
        per_doc.append(
            {
                "source": source,
                "page_count": ex.get("page_count"),
                "pages_analyzed": page_numbers,
            }
        )

        if not page_numbers:
            continue

        is_inss_doc = _is_inss_doc(ex)

        laudo_role = laudo_roles.get(source, "none")

        chunks = [page_numbers] if pages_per_call <= 0 else [page_numbers[i : i + pages_per_call] for i in range(0, len(page_numbers), pages_per_call)]
        for chunk_idx, chunk in enumerate(chunks):
            # Atualiza progresso se task foi fornecida
            if task:
                progress = 30 + int((ex_idx - 1) / len(extractions) * 65 + (chunk_idx / len(chunks)) * (65 / len(extractions)))
                task.update(progress=progress, message=f"Analisando {source} (parte {chunk_idx + 1}/{len(chunks)})...")
            
            chunk_text = build_text_from_ocr(ex, pages=chunk)
            if not chunk_text.strip():
                continue

            personal_data_policy = (
                "POLÍTICA DE DADOS PESSOAIS (OBRIGATÓRIA):\n"
                "- Os campos de 'qualificacao_parte_autora' representam dados pessoais do paciente/parte autora.\n"
            )
            if is_inss_doc:
                personal_data_policy += (
                    "- ESTE documento é o PDF do INSS.\n"
                    f"- Analise o conteúdo de todas as {len(page_numbers)} páginas e extraia 'qualificacao_parte_autora' de qualquer parte relevante do documento.\n"
                    "- ATENÇÃO: NUNCA use dados da seção 'Procuradores / Representantes Legais' como dados do paciente.\n"
                    "  - IGNORE 'Horlando Braga Filho' e seu CPF (028.951.113-59) - ele é o advogado/procurador, NÃO o paciente e NÃO é o representante legal.\n"
                    "\n"
                    "REGRAS PARA REPRESENTANTE LEGAL (CRÍTICO):\n"
                    "- Procure por seções com o título 'PROCURAÇÃO' ou texto contendo 'Outorgante' e 'representado (a) por'.\n"
                    "- FORMATO TÍPICO DA PROCURAÇÃO:\n"
                    "  'PROCURAÇÃO\n"
                    "   Outorgante: [Nome do Paciente], brasileiro, [estado civil], [profissão], CPF: XXX.XXX.XXX-XX e RG: XXXXXXXXX,\n"
                    "   representado (a) por [Nome do Representante Legal], CPF: XXX.XXX.XXX-XX e RG: XXXXXXXXX'\n"
                    "\n"
                    "- REGRA DE EXTRAÇÃO:\n"
                    "  1. OUTORGANTE = PACIENTE (pessoa que dá a procuração) → Preencha 'nome', 'cpf' da qualificacao_parte_autora\n"
                    "  2. REPRESENTADO POR = REPRESENTANTE LEGAL (quem representa o paciente) → Preencha:\n"
                    "     - representante_legal_nome: nome completo que vem APÓS 'representado (a) por' ou 'representado(a) por'\n"
                    "     - representante_legal_cpf: CPF que vem logo após o nome do representante\n"
                    "     - representante_legal_rg: RG que vem logo após o CPF do representante\n"
                    "\n"
                    "- EXEMPLO REAL:\n"
                    "  Texto: 'Outorgante: Heliabison Matias Correia, brasileiro, Solteiro(a), Estudante, CPF: 078.428.503-99\n"
                    "         e RG: 2020096219-6, representado (a) por Daiane Cunha Matias, CPF: 021.409.413-81 e RG: 2004003002105'\n"
                    "  EXTRAÇÃO CORRETA:\n"
                    "  - nome: 'Heliabison Matias Correia'\n"
                    "  - cpf: '078.428.503-99'\n"
                    "  - representante_legal_nome: 'Daiane Cunha Matias'\n"
                    "  - representante_legal_cpf: '021.409.413-81'\n"
                    "  - representante_legal_rg: '2004003002105'\n"
                    "\n"
                    "- ATENÇÃO: NÃO confunda o representante legal (familiar/tutor na procuração) com o advogado/procurador (Horlando Braga Filho).\n"
                )
            else:
                personal_data_policy += (
                    "- ESTE documento NÃO é o PDF do INSS.\n"
                    "- Portanto, retorne TODOS os campos de 'qualificacao_parte_autora' como null.\n"
                )

            laudo_policy = (
                "POLÍTICA DOS LAUDOS (OBRIGATÓRIA):\n"
                "- Existem 2 laudos médicos distintos. Não misture médico/especialidade/datas/CIDs entre eles.\n"
            )
            conclusao_policy = (
                "POLÍTICA DE CONCLUSÃO MÉDICA (OBRIGATÓRIA):\n"
                "- Extraia 'diagnostico_final_tratamento' (deficiência/CID, medicamento, finalidade). "
                "'conclusao_medica': síntese da conclusão clínica do médico. Preencha se houver nos laudos.\n"
            )
            relatorio_escolar_policy = (
                "POLÍTICA DO RELATÓRIO ESCOLAR (OBRIGATÓRIA):\n"
                "- Nos campos 'resumo' e 'resumo_continuacao' do relatorio_escolar: extraia SOMENTE o que é importante sobre o aluno "
                "(dificuldades, necessidades, limitações, recomendações, desempenho). "
                "NÃO inclua texto introdutório como 'este relatório visa fornecer informações...' ou similares.\n"
            )
            if laudo_role == "principal":
                laudo_policy += (
                    "- ESTE documento é o PRIMEIRO LAUDO (laudo_principal). Preencha COMPLETAMENTE 'dados_medicos.laudo_principal'.\n"
                    "- O campo 'trecho_clinico_relevante' (descrição do laudo) DEVE começar com 'Recomenda' (ex: 'Recomenda acompanhamento...', 'Recomenda-se que...').\n"
                    "- Retorne 'dados_medicos.laudo_psiquiatrico_segundo_laudo' como null.\n"
                )
            elif laudo_role == "segundo":
                laudo_policy += (
                    "- ESTE documento é o SEGUNDO LAUDO (laudo_psiquiatrico_segundo_laudo). Preencha COMPLETAMENTE 'dados_medicos.laudo_psiquiatrico_segundo_laudo'.\n"
                    "- Retorne 'dados_medicos.laudo_principal' como null.\n"
                )
            else:
                laudo_policy += (
                    "- Se este documento não for um dos laudos, retorne ambos os blocos de laudo como null.\n"
                )

            user_msg = (
                f"Documento fonte: {source}\n"
                "Analise o texto a seguir e preencha os campos do schema. Retorne somente o JSON.\n"
                + personal_data_policy
                + laudo_policy
                + conclusao_policy
                + relatorio_escolar_policy
                + "Texto:\n"
                + chunk_text
            )

            messages = [
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ]

            resp = call_openai_chat(model, messages, [FUNCTION_SCHEMA])
            chunk_structured = extract_from_response(resp)
            if is_inss_doc:
                _log_inss_extraction(source, chunk_idx, len(chunks), chunk_text, chunk_structured)
            if not structured_result:
                structured_result = chunk_structured
            else:
                merge_structured_field(structured_result, chunk_structured)

            if (chunk_idx < len(chunks) - 1) and delay_between_calls > 0:
                time.sleep(delay_between_calls)

        if (ex_idx < len(extractions)) and delay_between_calls > 0:
            time.sleep(delay_between_calls)

    return structured_result, per_doc


def generate_docx_from_structured(structured: Dict[str, Any], template_path: str) -> str:
    """Gera um DOCX preenchido e retorna o caminho do arquivo gerado (temp)."""
    replacements = prepare_replacements(structured)
    if os.path.exists(template_path):
        doc = Document(template_path)
        process_document(doc, replacements)
    else:
        doc = Document()
        doc.add_paragraph("Template não encontrado. Preencha/ajuste 'template_final_real.docx'.")
        doc.add_paragraph("Campos detectados:")
        for k, v in replacements.items():
            doc.add_paragraph(f"[{k}] = {v}")

    fd, out_path = tempfile.mkstemp(prefix="template_gerado_", suffix=".docx")
    os.close(fd)
    doc.save(out_path)
    return out_path


app = FastAPI(title="Extração de dados (PDF/IMG → JSON → DOCX)")

if os.path.isdir(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.on_event("startup")
def _startup() -> None:
    # carrega .env/.ENV se existir
    load_env(None)


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    index_path = os.path.join(STATIC_DIR, "index.html")
    if not os.path.exists(index_path):
        raise HTTPException(status_code=500, detail="static/index.html não encontrado.")
    with open(index_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/login", response_class=HTMLResponse)
def login_page() -> str:
    login_path = os.path.join(STATIC_DIR, "login.html")
    if not os.path.exists(login_path):
        raise HTTPException(status_code=500, detail="static/login.html não encontrado.")
    with open(login_path, "r", encoding="utf-8") as f:
        return f.read()


@app.post("/api/extract")
async def api_extract(
    _user: Dict[str, Any] = Depends(require_firebase_user),
    files: List[UploadFile] = File(...),
    model: str = DEFAULT_MODEL,
    pages_per_call: int = 0,
    delay_between_calls: float = 0.0,
    inss_max_pages: int = 6,
    inss_min_total_pages: int = 20,
    ocr_lang: str = "por",
    ocr_min_chars: int = 30,
    ocr_dpi: int = 300,
) -> Dict[str, Any]:
    _cleanup_store()
    _cleanup_tasks()

    if len(files) < 1:
        raise HTTPException(status_code=400, detail="Envie pelo menos 1 arquivo (PDF, PNG ou JPG).")

    # Valida e lê os arquivos
    file_contents: List[Tuple[bytes, str]] = []
    allowed_extensions = ('.pdf', '.png', '.jpg', '.jpeg')
    
    for f in files:
        filename_lower = (f.filename or "").lower()
        if not filename_lower.endswith(allowed_extensions):
            raise HTTPException(
                status_code=400, 
                detail=f"Arquivo deve ser PDF, PNG ou JPG: {f.filename}"
            )
        content = await f.read()
        if not content:
            raise HTTPException(status_code=400, detail=f"Arquivo vazio: {f.filename}")
        file_contents.append((content, f.filename or "documento.pdf"))

    # Cria uma tarefa e inicia processamento em background
    task_id = str(uuid.uuid4())
    task = Task(task_id)
    
    with _TASKS_LOCK:
        _TASKS[task_id] = task
    
    # Inicia thread de processamento
    thread = threading.Thread(
        target=_process_extraction_background,
        args=(
            task_id,
            file_contents,
            model,
            pages_per_call,
            delay_between_calls,
            inss_max_pages,
            inss_min_total_pages,
            ocr_lang,
            ocr_min_chars,
            ocr_dpi,
        ),
        daemon=True,
    )
    thread.start()
    
    return {"task_id": task_id, "message": "Processamento iniciado. Use /api/task/{task_id} para verificar o status."}


@app.get("/api/task/{task_id}")
def api_get_task_status(
    task_id: str,
    _user: Dict[str, Any] = Depends(require_firebase_user),
) -> Dict[str, Any]:
    """Retorna o status de uma tarefa em background."""
    _cleanup_tasks()
    
    with _TASKS_LOCK:
        task = _TASKS.get(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada ou expirada.")
    
    return task.to_dict()


@app.post("/api/generate-docx")
def api_generate_docx(
    payload: Dict[str, Any] = Body(...),
    _user: Dict[str, Any] = Depends(require_firebase_user),
) -> FileResponse:
    _cleanup_store()
    token = payload.get("token")
    if not token or not isinstance(token, str):
        raise HTTPException(status_code=400, detail="Campo obrigatório: token")
    item = _STORE.get(token)
    if not item:
        raise HTTPException(status_code=404, detail="Token não encontrado ou expirado. Refaça a extração.")

    structured_override = payload.get("structured")
    if structured_override is not None:
        if not isinstance(structured_override, dict):
            raise HTTPException(status_code=400, detail="Campo 'structured' deve ser um objeto JSON.")
        structured = structured_override
    else:
        structured = item.get("structured")
        if not isinstance(structured, dict):
            raise HTTPException(status_code=500, detail="Conteúdo estruturado inválido no servidor.")

    structured = _force_brazilian_nationality(structured)
    structured = _process_strings(structured)

    out_path = generate_docx_from_structured(structured, TEMPLATE_PATH)
    # opcional: invalidar token após uso (evita reuso)
    _STORE.pop(token, None)

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="template_gerado.docx",
    )

